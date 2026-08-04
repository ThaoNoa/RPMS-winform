# -*- coding: utf-8 -*-
"""Generate RPMS project overview Word document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(r"E:\DoAn\RPMS\Docs\TongQuanDuAn_RPMS.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2)

def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading_custom(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True,
                     color=RGBColor(0x1E, 0x3A, 0x5F) if level <= 2 else None)
    return p

def add_para(text, size=11, bold=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p

def add_bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=size)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, size=10, bold=True)
            # shade header
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "D6E3F0")
            hdr[i]._tePr = hdr[i]._tc.get_or_add_tcPr()
            hdr[i]._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9)
    doc.add_paragraph()
    return table

# ========== TITLE ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("TÀI LIỆU TỔNG QUAN DỰ ÁN")
set_run_font(r, size=20, bold=True, color=RGBColor(0x1E, 0x3A, 0x5F))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("RPMS — Rental Property Management System")
set_run_font(r, size=16, bold=True)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    "Phiên bản tài liệu: 2.0  |  .NET 8 WinForms + EF Core 8 + SQL Server\n"
    "Mục đích: Dev khác đọc để nắm toàn bộ source (layer, logic, biến, SQL, UI)"
)
set_run_font(r, size=10)

add_para("Ngày cập nhật: theo source hiện tại trong workspace E:\\DoAn\\RPMS", size=10)

# ========== TOC-like ==========
add_heading_custom("Mục lục nội dung", 1)
for item in [
    "1. Tổng quan & mục đích hệ thống",
    "2. Công nghệ & môi trường",
    "3. Kiến trúc phân lớp (Layers)",
    "4. Cấu trúc solution & dependency",
    "5. Database SQL (schema, constraint, sample)",
    "6. DAL — Entity, Repository, UnitOfWork, SchemaUpdater",
    "7. DTO — các nhóm dữ liệu truyền giữa layer",
    "8. Common — Constants, UserSession, AppLayout",
    "9. BLL — Services, Helpers, Mapping, Seeder, DI",
    "10. WinForms — Program, MainForm, Forms theo role",
    "11. Luồng nghiệp vụ chi tiết",
    "12. Quy ước đặt tên biến & trạng thái",
    "13. Tài khoản demo & cách chạy",
    "14. Phụ lục — checklist đọc source",
]:
    add_bullet(item)

# ========== 1 ==========
add_heading_custom("1. Tổng quan & mục đích hệ thống", 1)
add_para(
    "RPMS (Rental Property Management System) là ứng dụng Desktop quản lý cho thuê nhà/phòng trọ. "
    "Hệ thống phục vụ 4 vai trò cứng theo RoleID: Admin (1), Landlord/Chủ nhà (2), Tenant/Người thuê (3), "
    "Manager/Quản lý viên (4). Chức năng chính gồm: xác thực người dùng, quản lý nhà–phòng–tin đăng, "
    "lịch hẹn xem phòng, hợp đồng (nháp/Active, sửa chờ xác nhận), ghi chỉ số điện nước & hóa đơn "
    "(prorate giữa tháng / đổi giá), bảo trì, phân công Manager, chat, thông báo, dashboard & báo cáo."
)
add_para(
    "Nguyên tắc kiến trúc: UI (WinForms) chỉ gọi BLL qua interface; BLL dùng UnitOfWork/Repository; "
    "DAL dùng EF Core map tới SQL Server. Không gọi DbContext trực tiếp từ Form."
)

# ========== 2 ==========
add_heading_custom("2. Công nghệ & môi trường", 1)
add_table(
    ["Thành phần", "Chi tiết"],
    [
        ["Runtime", ".NET 8"],
        ["UI", "Windows Forms (net8.0-windows)"],
        ["ORM", "Entity Framework Core 8 + SQL Server provider"],
        ["Database", "SQL Server Express instance .\\SQLEXPRESS, DB name RPMS"],
        ["DI", "Microsoft.Extensions.DependencyInjection"],
        ["Mapping", "AutoMapper 12 (MappingProfile)"],
        ["Mật khẩu", "BCrypt.Net-Next (PasswordHelper)"],
        ["Auth session", "UserSession static (in-memory, không JWT)"],
    ],
)
add_para("Yêu cầu: Windows 10/11, .NET 8 SDK, SQL Server Express, Visual Studio 2022 (khuyến nghị).")

# ========== 3 ==========
add_heading_custom("3. Kiến trúc phân lớp (Layers)", 1)
add_code(
    "RPMS.WinForms  (UI + Program.cs DI host)\n"
    "      ↓ gọi interface\n"
    "RPMS.BLL       (Services, Helpers, DataSeeder, AutoMapper)\n"
    "      ↓ IUnitOfWork / Repositories\n"
    "RPMS.DAL       (EF Core DbContext, Entities, Configurations, Repos)\n"
    "      ↓\n"
    "SQL Server RPMS\n"
    "\n"
    "Ngang hàng: RPMS.DTO (DTO), RPMS.Common (Constants, UserSession)"
)
add_para(
    "Vòng đời request UI điển hình: Form → IXxxService (scoped/transient qua scope) → "
    "UnitOfWork.XxxRepository → RPMSContext → SQL → map Entity→DTO (AutoMapper hoặc Map thủ công) → Form bind DataGridView."
)

# ========== 4 ==========
add_heading_custom("4. Cấu trúc solution & dependency", 1)
add_table(
    ["Project", "TFM", "Vai trò", "Tham chiếu"],
    [
        ["RPMS.DTO", "net8.0", "DTO request/response", "—"],
        ["RPMS.Common", "net8.0", "AppColors, AppTypography, AppLayout, UserSession", "→ DTO"],
        ["RPMS.DAL", "net8.0", "EF, Entity, Repo, UoW, SchemaUpdater", "→ Common, DTO"],
        ["RPMS.BLL", "net8.0", "Business services", "→ DAL, DTO, Common"],
        ["RPMS.WinForms", "net8.0-windows", "WinExe UI", "→ BLL, Common, DTO"],
        ["BCryptHelper", "net8.0 Exe", "Tool hash mật khẩu standalone", "BCrypt"],
    ],
)
add_para("Thư mục khác: Database/ (RPMS_Full.sql), Docs/, tools/ (smoke/probe), README.md.")

# ========== 5 ==========
add_heading_custom("5. Database SQL", 1)
add_heading_custom("5.1. Script chính", 2)
add_para(
    "File: Database/RPMS_Full.sql — DROP/CREATE database RPMS, tạo toàn bộ bảng + index + CHECK + sample data. "
    "Đường dẫn FILENAME .mdf/.ldf trong script có thể cần sửa theo máy. Chạy khuyến nghị:"
)
add_code("sqlcmd -S .\\SQLEXPRESS -E -f 65001 -i Database\\RPMS_Full.sql")

add_heading_custom("5.2. Danh sách bảng & khóa", 2)
add_table(
    ["Bảng", "PK", "FK / Unique chính", "Status / CHECK nổi bật"],
    [
        ["Roles", "RoleID", "RoleName UNIQUE", "—"],
        ["Users", "UserID", "FK RoleID; Username/Email UNIQUE", "Active|Inactive"],
        ["Houses", "HouseID", "FK OwnerID→Users", "Active|Inactive"],
        ["Rooms", "RoomID", "FK HouseID; UQ(HouseID,RoomNumber)", "Available|Occupied|Maintenance"],
        ["RoomImages", "ImageID", "FK RoomID CASCADE", "—"],
        ["Amenities", "AmenityID", "AmenityName UNIQUE", "—"],
        ["RoomAmenities", "RoomAmenityID", "UQ(RoomID,AmenityID)", "—"],
        ["Posts", "PostID", "FK RoomID", "Pending|Approved|Rejected|Expired|Hidden"],
        ["PostImages", "PostImageID", "FK PostID", "—"],
        ["Favorites", "FavoriteID", "UQ(UserID,RoomID)", "—"],
        ["Appointments", "AppointmentID", "FK Room, Tenant", "Pending|Accepted|Rejected|Completed (+Cancelled ở app)"],
        ["Contracts", "ContractID", "FK Room, Tenant?, CreatedBy", "Draft|Active|Expired|Terminated"],
        ["MeterReadings", "ReadingID", "FK Contract, CreatedBy", "New≥Old"],
        ["Invoices", "InvoiceID", "FK Contract, Reading", "Unpaid|Paid|Overdue|Cancelled"],
        ["Payments", "PaymentID", "FK Invoice", "Completed|Failed|Refunded"],
        ["MaintenanceRequests", "RequestID", "FK Contract, Manager?", "Pending|Processing|Completed|Cancelled"],
        ["Assignments", "AssignmentID", "UQ(HouseID,ManagerID)", "Active|Inactive"],
        ["Reviews", "ReviewID", "FK Contract (1-1)", "Rating 1–5"],
        ["Notifications", "NotificationID", "FK User", "IsRead bit"],
        ["ActivityLogs", "LogID", "FK User", "—"],
    ],
)
add_para(
    "Bảng Chat (ChatConversations, ChatMessages) có thể không nằm trong script SQL gốc — "
    "được tạo runtime bởi DatabaseSchemaUpdater."
)

add_heading_custom("5.3. Cột Contracts mở rộng (schema updater)", 2)
add_para(
    "PendingMonthlyRent, PendingElectricPrice, PendingWaterPrice, PendingDeposit, PendingEndDate, "
    "PendingEditStatus, PendingEditNote, PendingEditAt, PreviousMonthlyRent, PreviousElectricPrice, "
    "PreviousWaterPrice, PriceEffectiveDate. TenantID cho phép NULL (hợp đồng nháp)."
)

add_heading_custom("5.4. Sample Roles & Users", 2)
add_table(
    ["RoleID", "RoleName", "User mẫu (Username / Password plain trước hash)"],
    [
        ["1", "Admin", "admin / admin123"],
        ["2", "Landlord", "namlandlord / 123456"],
        ["3", "Tenant", "tenant / 123456"],
        ["4", "Manager", "manager / 123456"],
    ],
)
add_para(
    "DataSeeder khi mở app: hash mật khẩu plain → BCrypt; sửa tên tiếng Việt nếu mojibake; "
    "đồng bộ timeline hợp đồng/hóa đơn sample theo tháng hiện tại."
)

# ========== 6 ==========
add_heading_custom("6. DAL — Entity, Repository, UnitOfWork", 1)
add_heading_custom("6.1. RPMSContext", 2)
add_para(
    "File: RPMS.DAL/Data/RPMSContext.cs — DbSet cho mọi entity. "
    "Cấu hình Fluent API trong RPMS.DAL/Configurations/*Configuration.cs "
    "(ToTable, HasKey, MaxLength, CheckConstraint, quan hệ Restrict/Cascade)."
)

add_heading_custom("6.2. Entity — thuộc tính quan trọng", 2)
entities = [
    ("User", "UserID, RoleID, FullName, Phone, Email, Username, Password, Address, Status, CreatedDate, UpdatedDate; nav Role, Houses, Assignments, Notifications…"),
    ("House", "HouseID, OwnerID, HouseName, Address, Description, Status; nav Owner, Rooms, Assignments"),
    ("Room", "RoomID, HouseID, RoomNumber, Floor, Area, Price, Capacity, Bedroom, Bathroom, Furniture, Status; nav House, Images, Amenities, Contracts, Posts"),
    ("Contract", "ContractID, ContractCode, RoomID, TenantID?, StartDate, EndDate, MoveInDate, MoveOutDate?, Deposit, MonthlyRent, ElectricPrice, WaterPrice, Status, CreatedBy + pending/previous price fields"),
    ("Invoice", "InvoiceID, InvoiceCode, ContractID, ReadingID, Rent, ElectricCost, WaterCost, OtherFee, Total, Status, DueDate, PaidDate"),
    ("MeterReading", "ReadingID, ContractID, ReadingMonth, OldElectric, NewElectric, OldWater, NewWater, CreatedBy"),
    ("Appointment", "AppointmentID, RoomID, TenantID, AppointmentDate, Status, Note"),
    ("Assignment", "AssignmentID, HouseID, ManagerID, AssignedDate, Status"),
    ("MaintenanceRequest", "RequestID, ContractID, Title, Description, Image, Status, AssignedManager, CompletedDate"),
    ("Notification", "NotificationID, UserID, Title, Content, IsRead, CreatedDate, UpdatedDate"),
    ("ChatConversation / ChatMessage", "LandlordID+TenantID pair; Message: SenderID, Content, ImagePath, IsRead"),
]
for name, props in entities:
    add_bullet(f"{name}: {props}")

add_heading_custom("6.3. GenericRepository & UnitOfWork", 2)
add_para(
    "IGenericRepository<T>: GetAllAsync(include), GetByIdAsync, FindAsync(expression, include), "
    "FirstOrDefaultAsync, Add/AddRange, Update, Remove, ExistsAsync, CountAsync. "
    "Include properties dạng chuỗi \"Room.House, Tenant\"."
)
add_para(
    "IUnitOfWork expose: Roles, Users, Houses, Rooms, RoomImages, Amenities, RoomAmenities, Posts, "
    "PostImages, Favorites, Appointments, Contracts, Reviews, MeterReadings, Invoices, Payments, "
    "MaintenanceRequests, Assignments, Notifications, ActivityLogs, ChatConversations, ChatMessages; "
    "SaveChangesAsync; BeginTransactionAsync / Commit / Rollback."
)

add_heading_custom("6.4. DatabaseSchemaUpdater", 2)
add_para(
    "EnsureUpdatedAsync: EnsureCreatedAsync; patch Reviews reply columns; tạo Chat tables; "
    "nullable TenantID; CHECK Draft; thêm từng cột pending/previous trên Contracts (tách ExecAsync "
    "để tránh rollback cả batch)."
)

add_heading_custom("6.5. DI DAL", 2)
add_code(
    "services.AddDbContext<RPMSContext>(…UseSqlServer(cs));\n"
    "services.AddScoped(typeof(IGenericRepository<>), typeof(GenericRepository<>));\n"
    "services.AddScoped<IUserRepository, UserRepository>(); // … mọi repo\n"
    "services.AddScoped<IUnitOfWork, UnitOfWork>();"
)

# ========== 7 ==========
add_heading_custom("7. DTO — nhóm dữ liệu", 1)
add_para("Thư mục RPMS.DTO theo domain. Một số DTO tiêu biểu:")
add_table(
    ["Nhóm", "DTO chính", "Ghi chú"],
    [
        ["Auth", "LoginRequestDto, LoginResponseDto, ChangePasswordDto", "LoginResponse: UserID, RoleID, RoleName, FullName, Username…"],
        ["User", "UserDto, CreateUserDto, UpdateUserDto", "RoleName map từ Role"],
        ["House/Room", "HouseDto, RoomDto, RoomDetailDto, Create*", "TotalRooms, Images, Amenities"],
        ["Post", "PostDto, PostDetailDto, RoomSearchFilterDto, CreatePostDto", "Bộ lọc tìm phòng nâng cao"],
        ["Contract", "ContractDto, ContractDetailDto, CreateContractDto, AssignTenantDto, UpdateContractDto, BulkCreateDraftContractsDto", "PendingEditStatus trên list"],
        ["Invoice", "InvoiceDto, InvoiceDetailDto, GenerateInvoiceDto, MeterReadingSummaryDto", "CreatedBy trên Generate"],
        ["Assignment", "AssignmentDto, CreateAssignmentDto", "HouseID + ManagerID"],
        ["Tenant", "AppointmentDto, TenantDashboardDto…", "RoomNumber, TenantName"],
        ["Maintenance", "MaintenanceRequestDto, Create*", "Timeline status"],
        ["Notification", "NotificationDto, CreateNotificationDto", "IsRead"],
        ["Statistic/Report", "Admin/Landlord/ManagerDashboardDto, Report*", "RevenueByMonth, Occupancy"],
        ["Chat/Calendar", "ConversationDto, ChatMessageDto, CalendarEventDto", "—"],
    ],
)

# ========== 8 ==========
add_heading_custom("8. Common", 1)
add_bullet("AppColors: Primary #2563EB, Background #F8FAFC, Card #FFFFFF, Success/Warning/Danger, Sidebar #0F172A, TextMain/Muted…")
add_bullet("AppTypography: Title/Subtitle/Heading/Body/BodyBold/Caption/Button — Clone() Font Segoe UI an toàn WinForms.")
add_bullet("AppLayout: PageHeaderHeight=56, ToolbarHeight=72, PagePadding=16, ButtonHeight=40, SidePanelWidth=360, DialogMin, PageMin…")
add_bullet("UserSession: static CurrentUser (LoginResponseDto?), Login(), Logout(), IsLoggedIn — không persist sau khi tắt app.")

# ========== 9 ==========
add_heading_custom("9. BLL — Services, Helpers, Mapping, Seeder", 1)
add_heading_custom("9.1. Đăng ký DI (BllDependencyInjection)", 2)
add_para(
    "AddAutoMapper(assembly); AddScoped cho: Auth, User, Role, House, Room, Amenity, Post, Contract, "
    "Invoice, Maintenance, Statistic, TenantInteraction, Landlord, Tenant, Notification, Assignment, "
    "ActivityLog, Review, Chat, Calendar, Report. Backup đăng ký Singleton riêng trong Program nếu có."
)

add_heading_custom("9.2. Catalog service (phương thức chính)", 2)
services_doc = [
    ("AuthService", "LoginAsync, ChangePasswordAsync, ResetPasswordAsync"),
    ("UserService", "GetAll/ByRole/ById, Create, Update, Delete, ToggleUserStatusAsync"),
    ("HouseService", "CRUD nhà theo OwnerID"),
    ("RoomService", "CRUD phòng, UpdateStatus, UploadRoomImagesAsync, AssignAmenitiesAsync"),
    ("PostService", "Create, GetActive/Pending, Approve/Reject, IncrementViewCount"),
    ("ContractService", "Get theo Tenant/Landlord/Manager; Create; CreateDraftContractsForHouseAsync; AssignTenant; Update (Pending); Confirm/Reject/CancelPending edit; Terminate; Extend"),
    ("InvoiceService", "GetByContract/ById; GetLatestReading; GenerateMonthlyInvoiceAsync; ProcessPaymentAsync"),
    ("LandlordService", "GetAppointments; GetAppointmentTenants; UpdateAppointmentStatusAsync (+notify); CreateNotificationForTenants"),
    ("TenantInteractionService", "BookAppointmentAsync; ToggleFavorite; Get/Remove Favorites"),
    ("TenantService", "GetTenantDashboardAsync; SearchRoomsAsync; SendContractRequestAsync"),
    ("AssignmentService", "GetAll/ByLandlord/ByManager; CreateAsync(landlordId); DeactivateAsync"),
    ("MaintenanceService", "Create; Get theo house/tenant/manager; UpdateStatus; SendMaintenanceNotification"),
    ("NotificationService", "GetByUser; UnreadCount; MarkRead/All; Delete; Create"),
    ("StatisticService", "GetAdmin/Landlord/ManagerDashboardStatsAsync"),
    ("ReviewService", "CreateReview; Reply; GetByLandlord/Tenant/All; AverageRating"),
    ("ChatService", "GetConversations; GetOrCreate; GetMessages; Send; MarkRead; UnreadCount"),
    ("CalendarService", "GetEventsAsync(user, from, to)"),
    ("ReportService", "GetAdminReportAsync; GetLandlordReportAsync"),
    ("ActivityLogService", "LogAsync; GetRecent; GetByUser"),
]
for name, methods in services_doc:
    add_bullet(f"{name}: {methods}")

add_heading_custom("9.3. Helpers", 2)
add_bullet("PasswordHelper: HashPassword(plain), VerifyPassword(plain, hash) — BCrypt.")
add_bullet("RentProrationHelper.Calculate(monthlyRent, monthStart, contractStart, contractEnd, moveIn, moveOut) → OccupiedDays, DaysInMonth, ProratedRent, Note.")
add_bullet("ContractPricingHelper.WeightedUnitCost / CalculateRent — tính điện/nước/tiền nhà khi có Previous* + PriceEffectiveDate giữa tháng.")
add_bullet("Exceptions: NotFoundException, BadRequestException, RPMSException — Form bắt BadRequest → ShowWarning.")

add_heading_custom("9.4. MappingProfile (AutoMapper)", 2)
add_para(
    "CreateMap Entity→Dto với ForMember cho tên quan hệ (OwnerName, RoomNumber, TenantName, "
    "ảnh chính Post, Amenities…). ContractDto map PendingEditStatus, TenantName \"(Chưa có khách)\" nếu null."
)

add_heading_custom("9.5. DataSeeder", 2)
add_para(
    "SeedAsync(db): HashIfPlaintext cho admin/namlandlord/tenant/manager; FixSampleDisplayNames; "
    "SyncSampleTimeline (hợp đồng, meter, invoice, appointment theo DateTime.Now); seed Roles/Users nếu DB trống."
)

# ========== 10 ==========
add_heading_custom("10. WinForms — Program, MainForm, Forms", 1)
add_heading_custom("10.1. Program.cs", 2)
add_code(
    "ConnectionString = Server=.\\SQLEXPRESS;Database=RPMS;Trusted_Connection=True;\n"
    "TrustServerCertificate=True;MultipleActiveResultSets=True;\n\n"
    "ConfigureServices → AddDataAccessLayer + AddBusinessLogicLayer + Transient Forms\n"
    "Startup: SchemaUpdater.EnsureUpdatedAsync → DataSeeder.SeedAsync\n"
    "Loop: LoginForm.ShowDialog OK → MainForm.ShowDialog (Retry = đăng xuất)"
)
add_para(
    "Lưu ý DI: Form Transient resolve từ root ServiceProvider; service Scoped có thể bị giữ lâu trên root. "
    "Các form quan trọng (Contract, Assignment, Appointment, Meter, Notification, Chat) nên dùng "
    "IServiceScopeFactory.CreateScope() mỗi thao tác để tránh DbContext concurrency/stale."
)

add_heading_custom("10.2. MainForm — menu theo RoleID", 2)
add_para("Chung: Dashboard, Notifications, Profile, Calendar. Role 1|2: Reports. Role 2|3: Chat.")
add_table(
    ["RoleID", "Tag menu", "Form"],
    [
        ["1", "UserManagement, PostManagement, AdminReviews, ActivityLog, Backup", "Admin/*"],
        ["2", "LandlordHouse, LandlordRoom, LandlordAssignment, LandlordContract, LandlordAppointment, LandlordPost, LandlordReviews", "Landlord/*"],
        ["3", "TenantHome, TenantFavorite, TenantContract, TenantInvoice, TenantMaintenance, TenantReviews", "Tenant/*"],
        ["4", "ManagerMeter, ManagerMaintenance", "Manager/*"],
    ],
)
add_para(
    "LoadChildForm(tag): GetRequiredService<form> → OpenChildForm (TopLevel=false, Dock Fill vào pnlContent). "
    "Lỗi ctor/UI → \"Không mở được màn hình: …\"."
)

add_heading_custom("10.3. UI chuẩn (UIHelper)", 2)
add_bullet("CreatePageHeader(title, trailingButtons…) — TableLayout tiêu đề + nút phải; Tag = Label tiêu đề; GetPageHeaderTitle(panel).")
add_bullet("CreateFilterBar / CreateLabeledField — FlowLayout wrap, tránh absolute X đè nhau.")
add_bullet("CreateSideFormPanel / CreateDialogFooter / WireListPage / WirePage / ApplyGridFill.")
add_bullet("Controls: ModernButton, ModernTextBox (placeholder), ModernDataGridView, EmptyStatePanel, LoadingPanel, ToastNotifier, SidebarButton, SummaryCard, OccupancyChartPanel, StatusTimelineControl, RoomCardControl.")

add_heading_custom("10.4. Danh sách Form theo thư mục", 2)
forms = [
    ("Auth/LoginForm", "Đăng nhập → UserSession"),
    ("Auth/RegisterForm", "Đăng ký user"),
    ("Layout/MainForm", "Shell sidebar + host"),
    ("Dashboard/DashboardForm", "Thống kê theo role"),
    ("Admin/UserManagementForm + UserModalForm", "CRUD user"),
    ("Admin/PostManagementForm + PostDetailModalForm", "Duyệt tin"),
    ("Admin/ReviewManagementForm, ActivityLogForm, BackupForm", "Review / log / backup"),
    ("Landlord/House|Room Form+Modal", "CRUD nhà phòng media tiện nghi"),
    ("Landlord/LandlordAssignmentForm", "Gán Manager theo ID/Username"),
    ("Landlord/LandlordContractForm", "HĐ tạo/bulk draft/gán khách/sửa"),
    ("Landlord/LandlordAppointmentForm", "Xác nhận lịch + notify"),
    ("Landlord/LandlordPostForm, LandlordReviewForm", "Tin đăng / trả lời review"),
    ("Tenant/TenantHomeForm, RoomDetailForm", "Tìm phòng + gallery"),
    ("Tenant/TenantAppointmentModalForm, TenantFavoriteForm", "Đặt lịch / yêu thích"),
    ("Tenant/TenantContractForm, TenantInvoiceForm, InvoiceDetailForm", "HĐ confirm edit / hóa đơn TT"),
    ("Tenant/TenantMaintenanceForm, TenantReviewForm", "Báo sự cố / đánh giá"),
    ("Manager/ManagerMeterForm", "Ghi điện nước + tạo hóa đơn tháng trước"),
    ("Manager/ManagerMaintenanceForm, MaintenanceDetailForm", "Xử lý sự cố + timeline"),
    ("Shared/NotificationCenterForm, ProfileForm, ChatForm, CalendarForm, ReportForm", "Dùng chung"),
]
for f, d in forms:
    add_bullet(f"{f}: {d}")

# ========== 11 ==========
add_heading_custom("11. Luồng nghiệp vụ chi tiết", 1)

add_heading_custom("11.1. Đăng nhập", 2)
add_para(
    "LoginForm → AuthService.LoginAsync(username, password) → tìm User include Role → Verify BCrypt → "
    "Status Active → ActivityLog → LoginResponseDto → UserSession.Login → MainForm.GenerateMenu(RoleID)."
)

add_heading_custom("11.2. Hợp đồng", 2)
add_bullet("CreateContractAsync: phòng không Occupied; không HĐ Active/Draft; TenantID null → Status=Draft; có tenant → Active + Room Occupied + Notification.")
add_bullet("CreateDraftContractsForHouseAsync: mọi phòng nhà của landlord chưa có HĐ Active/Draft, không Occupied/Inactive → Draft; MonthlyRent = form hoặc Room.Price.")
add_bullet("AssignTenantAsync: gán khách cho Draft → Active + Occupied + notify.")
add_bullet("UpdateContractAsync: Draft/không tenant áp dụng ngay; Active+tenant → Pending* + notify; Confirm lưu Previous* + PriceEffectiveDate; Reject/CancelPending.")
add_bullet("GetContractsByManagerAsync: Assignments Active của manager → HouseIDs → Contracts của Room thuộc nhà đó.")

add_heading_custom("11.3. Ghi chỉ số & hóa đơn", 2)
add_para(
    "ManagerMeterForm lọc HĐ Active có TenantID. BillingMonth = tháng trước (đã kết thúc). "
    "GenerateMonthlyInvoiceAsync: không trùng tháng; Old từ reading trước; "
    "WeightedUnitCost điện/nước; CalculateRent + prorate ngày ở; tạo MeterReading + Invoice Unpaid + notify tenant. "
    "ProcessPaymentAsync: Paid + Payment record."
)

add_heading_custom("11.4. Lịch hẹn", 2)
add_para(
    "Tenant BookAppointmentAsync → notify Landlord. Landlord UpdateAppointmentStatusAsync(Accepted/Rejected/Completed…) "
    "→ notify Tenant (tiếng Việt, kèm phòng/nhà/giờ). GetAppointmentTenantsAsync: tenant từng đặt lịch (không Rejected/Cancelled) để gán HĐ."
)

add_heading_custom("11.5. Phân công Manager", 2)
add_para(
    "LandlordAssignmentForm: tìm Manager (UserID hoặc Username, RoleID=4, Active) → CreateAsync(houseId, managerId, landlordId) "
    "kiểm tra nhà thuộc landlord; unique (House,Manager) → insert hoặc reactivate; notify Manager."
)

add_heading_custom("11.6. Tin đăng & tìm phòng", 2)
add_para(
    "Landlord tạo Post (Pending) → Admin Approve/Reject. Tenant SearchRoomsAsync / TenantHomeForm lọc giá, diện tích, "
    "amenities, status, featured; RoomDetailForm gallery ảnh/video."
)

# ========== 12 ==========
add_heading_custom("12. Quy ước đặt tên biến & trạng thái", 1)
add_heading_custom("12.1. Tên biến / ID", 2)
add_table(
    ["Kiểu", "Quy ước", "Ví dụ"],
    [
        ["Khóa chính", "EntityName + ID", "UserID, HouseID, ContractID"],
        ["Khóa ngoại", "Tên quan hệ + ID", "OwnerID, TenantID, ManagerID, CreatedBy"],
        ["DTO", "…Dto / Create…Dto / Update…Dto", "CreateContractDto"],
        ["Service", "IXxxService / XxxService", "IContractService"],
        ["Repo", "IXxxRepository / XxxRepository", "IContractRepository"],
        ["Form field", "cbo/txt/dtp/dgv/btn/pnl/lbl + Name", "cboHouse, dgvContracts"],
        ["Session", "UserSession.CurrentUser", "UserID, RoleID"],
    ],
)

add_heading_custom("12.2. Giá trị Status thường dùng (chuỗi)", 2)
add_table(
    ["Đối tượng", "Giá trị"],
    [
        ["User / House / Assignment", "Active, Inactive"],
        ["Room", "Available, Occupied, Maintenance"],
        ["Post", "Pending, Approved, Rejected, Expired, Hidden"],
        ["Appointment", "Pending, Accepted, Rejected, Completed, Cancelled"],
        ["Contract", "Draft, Active, Expired, Terminated"],
        ["PendingEditStatus", "Pending (và null khi không chờ)"],
        ["Invoice", "Unpaid, Paid, Overdue, Cancelled"],
        ["Maintenance", "Pending, Processing, Completed, Cancelled"],
    ],
)

add_heading_custom("12.3. Mã sinh tự động", 2)
add_bullet("ContractCode: HD + yyyyMMddHHmmss + RoomID (4 số) [+ seq bulk]")
add_bullet("InvoiceCode: INV + yyMMddHHmmss")

# ========== 13 ==========
add_heading_custom("13. Tài khoản demo & cách chạy", 1)
add_code(
    "dotnet restore RPMS.sln\n"
    "dotnet build RPMS.sln\n"
    "dotnet run --project RPMS.WinForms"
)
add_table(
    ["Username", "Password", "Role"],
    [
        ["admin", "admin123", "Admin"],
        ["namlandlord", "123456", "Landlord"],
        ["tenant", "123456", "Tenant"],
        ["manager", "123456", "Manager"],
    ],
)
add_para(
    "Đổi connection string: RPMS.WinForms/Program.cs property ConnectionString. "
    "Đóng process RPMS.WinForms trước khi build nếu bị khóa DLL."
)

# ========== 14 ==========
add_heading_custom("14. Phụ lục — checklist đọc source", 1)
for tip in [
    "Bắt đầu Program.cs → BllDependencyInjection → DalDependencyInjection → MainForm.GenerateMenu.",
    "Muốn hiểu HĐ: ContractService + LandlordContractForm + TenantContractForm + entity Contract.",
    "Muốn hiểu hóa đơn: InvoiceService + ManagerMeterForm + ContractPricingHelper + RentProrationHelper.",
    "Muốn hiểu phân quyền: RoleID trong Users + switch MainForm + kiểm tra OwnerID/ManagerID trong service.",
    "Muốn hiểu DB: Database/RPMS_Full.sql rồi DatabaseSchemaUpdater (patch runtime).",
    "UI layout: UIHelper + AppLayout; tránh absolute X cho toolbar.",
    "DbContext: ưu tiên IServiceScopeFactory trong Form async UI.",
    "Mapping: MappingProfile.cs + DTO cùng tên property entity.",
]:
    add_bullet(tip)

add_para(
    "— Hết tài liệu tổng quan RPMS v2.0. Cập nhật tài liệu khi thêm service/bảng/form mới. —",
    size=10, bold=True
)

doc.save(OUT)
print("Wrote", OUT)
print("Size", OUT.stat().st_size, "bytes")
